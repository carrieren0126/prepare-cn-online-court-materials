#!/usr/bin/env python3
"""Build draft/final Chinese court or police-report material packets from JSON.

The script is deterministic packaging only. Legal characterization and factual
normalization must happen before invocation. It never signs or submits files.
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageOps
from pypdf import PdfReader, PdfWriter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    LongTable,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


FONT_CN = "CourtPacketCJK"
FONT_DOC = "Heiti SC"


def register_cjk_font() -> str:
    candidates = [
        os.environ.get("CN_COURT_FONT"),
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            pdfmetrics.registerFont(TTFont(FONT_CN, candidate, subfontIndex=0))
            return candidate
    raise RuntimeError(
        "No embeddable Chinese font found. Set CN_COURT_FONT to a CJK TTF/TTC file."
    )


EMBEDDED_CJK_FONT = register_cjk_font()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("case_json", type=Path)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--mode", choices=("draft", "final"), default="draft")
    return parser.parse_args()


def load_case(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        raise ValueError("case JSON must contain an object")
    data["_base_dir"] = str(path.resolve().parent)
    return data


def text(value: Any, fallback: str = "待确认") -> str:
    if value is None or value == "":
        return fallback
    return str(value)


def critical_issues(data: dict[str, Any]) -> list[str]:
    issues: list[str] = []
    route = data.get("route")
    if route not in {"civil", "police", "review"}:
        issues.append("处理路径必须是民事起诉、报案或人工核验之一")
    if route == "review":
        issues.append("当前为人工核验路径，不能生成正式提交文书")
    plaintiff = data.get("plaintiff") or {}
    if not plaintiff.get("name"):
        issues.append("原告或报案人姓名缺失")
    if not data.get("facts"):
        issues.append("基本事实缺失")
    for fact in data.get("facts") or []:
        if fact.get("confidence") in {"inferred", "missing"}:
            issues.append(f"事实尚未确认：{fact.get('text', '未命名事实')}")
    if route == "civil":
        if data.get("case_type") != "online_goods":
            issues.append("第一版仅支持个人间网络商品交易生成民事起诉状；其他类型须转报案或人工核验路径")
        transaction = data.get("transaction") or {}
        if transaction.get("party_relationship") != "individual_to_individual":
            issues.append("尚未确认交易双方均为个人，不能生成第一版正式起诉状")
        if not data.get("claims"):
            issues.append("具体诉讼请求缺失")
        defendants = data.get("defendants") or []
        if not defendants:
            issues.append("被告信息缺失")
        for defendant in defendants:
            if not defendant.get("name"):
                issues.append("被告姓名或单位名称缺失")
            if not defendant.get("distinguishable_basis"):
                issues.append("能够将被告与他人区分的身份线索不足")
        court = data.get("court") or {}
        if not court.get("name") or "待" in str(court.get("name")):
            issues.append("受诉法院尚未确认")
        if not court.get("live_upload_rules_checked"):
            issues.append("尚未核对人民法院在线服务当前上传要求")
    for item in data.get("evidence") or []:
        if not item.get("source"):
            issues.append(f"证据来源缺失：{item.get('id', '?')}")
        if not item.get("proof_purpose"):
            issues.append(f"证据证明目的缺失：{item.get('id', '?')}")
        if not item.get("confirmed", False):
            issues.append(f"证据尚未确认：{item.get('id', '?')}")
    signature = data.get("signature") or {}
    if not signature.get("name") or "待" in str(signature.get("name")):
        issues.append("尚待本人签名或盖章")
    if not signature.get("date") or "待" in str(signature.get("date")):
        issues.append("提交日期尚未确认")
    for issue in data.get("open_issues") or []:
        if issue.get("severity") == "critical":
            issues.append(text(issue.get("item"), "存在关键待确认事项"))
    return list(dict.fromkeys(issues))


def validate(data: dict[str, Any], mode: str) -> list[str]:
    issues = critical_issues(data)
    if data.get("route") == "civil" and (
        data.get("case_type") != "online_goods"
        or (data.get("transaction") or {}).get("party_relationship") != "individual_to_individual"
    ):
        raise ValueError("不支持生成起诉状：\n- " + "\n- ".join(issues))
    if mode == "final" and issues:
        raise ValueError("正式版本已被阻止：\n- " + "\n- ".join(issues))
    return issues


def set_run_font(run, name: str = FONT_DOC, size: float = 11, bold: bool = False) -> None:
    # Use an installed macOS CJK font for deterministic LibreOffice rendering.
    # Map conventional legal-document font requests to the available family.
    actual_name = FONT_DOC if name in {"宋体", "黑体", "SimSun", "SimHei"} else name
    run.font.name = actual_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), actual_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), actual_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), actual_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), "777777")


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9)


def init_doc(title: str, draft: bool) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_DOC
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_DOC)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_DOC)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_DOC)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(("草稿-待核对\n" if draft else "") + title)
    set_run_font(r, name="黑体", size=18, bold=True)
    if draft:
        r.font.color.rgb = None
    add_page_field(section.footer.paragraphs[0])
    return doc


def add_heading(doc: Document, heading: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(heading)
    set_run_font(r, name="黑体", size=13, bold=True)


def add_kv_table(doc: Document, rows: Iterable[tuple[str, str]]) -> None:
    rows = list(rows)
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.0)
    table.columns[1].width = Cm(12.0)
    set_table_borders(table)
    for idx, (label, value) in enumerate(rows):
        for cell in table.rows[idx].cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(idx, 0), "F2F2F2")
        p0 = table.cell(idx, 0).paragraphs[0]
        p1 = table.cell(idx, 1).paragraphs[0]
        set_run_font(p0.add_run(label), name="黑体", size=10.5, bold=True)
        set_run_font(p1.add_run(value), size=10.5)


def person_summary(person: dict[str, Any]) -> str:
    parts = [text(person.get("name"))]
    labeled_keys = (
        ("gender", "性别"),
        ("birth_date", "出生日期"),
        ("id_number", "公民身份号码"),
        ("phone", "联系电话"),
        ("registered_address", "户籍地址"),
        ("usual_address", "经常居住地"),
        ("address", "联系地址"),
        ("platform_account", "平台账号"),
        ("payment_account", "支付账号"),
    )
    for key, label in labeled_keys:
        value = person.get(key)
        if value and not str(value).startswith("待确认"):
            parts.append(f"{label}：{value}")
    return "；".join(parts)


def add_signature(doc: Document, data: dict[str, Any], role: str) -> None:
    signature = data.get("signature") or {}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    set_run_font(
        p.add_run(
            f"{role}（本人签名）：________________    "
            f"姓名：{text(signature.get('name'))}    日期：{text(signature.get('date'))}"
        ),
        size=10.5,
    )


def build_readiness_docx(data: dict[str, Any], issues: list[str], path: Path, draft: bool) -> None:
    doc = init_doc("材料准备情况及待确认事项", draft)
    add_heading(doc, "当前建议路径")
    route_labels = {"civil": "优先准备民事起诉材料", "police": "优先准备报案与证据保全材料", "review": "暂缓形成最终文书，先人工核验"}
    doc.add_paragraph(route_labels.get(data.get("route"), "待确认"))
    add_heading(doc, "仍需处理")
    if issues:
        for idx, issue in enumerate(issues, 1):
            p = doc.add_paragraph()
            set_run_font(p.add_run(f"{idx}. {issue}"))
    else:
        doc.add_paragraph("未发现阻止生成最终版本的关键缺项；仍须由本人逐项核对。")
    add_heading(doc, "重要提示")
    for item in (
        "本材料由现有资料整理生成，不代表法院或公安机关将受理、立案或支持相关请求。",
        "请保留原始设备、原始文件及平台可下载记录；不要修改唯一原件。",
        "正式提交前请核对身份、金额、日期、法院、请求、证据编号及签名。",
    ):
        doc.add_paragraph(item)
    doc.save(path)


def build_primary_docx(data: dict[str, Any], path: Path, draft: bool) -> None:
    route = data.get("route")
    title = "民事起诉状" if route == "civil" else "报案事实说明"
    doc = init_doc(title, draft)
    plaintiff = data.get("plaintiff") or {}
    role = "原告" if route == "civil" else "报案人"
    add_heading(doc, "当事人信息")
    rows = [(role, person_summary(plaintiff))]
    for idx, defendant in enumerate(data.get("defendants") or [], 1):
        rows.append(("被告" if route == "civil" and idx == 1 else f"被告{idx}" if route == "civil" else f"对方/线索{idx}", person_summary(defendant)))
    add_kv_table(doc, rows)
    if route == "civil":
        add_heading(doc, "受诉法院及案由")
        court = data.get("court") or {}
        add_kv_table(doc, [("受诉法院", text(court.get("name"))), ("管辖核验", text(court.get("basis"))), ("案由建议", text(data.get("title")))])
        add_heading(doc, "诉讼请求")
        for idx, claim in enumerate(data.get("claims") or [], 1):
            doc.add_paragraph(f"{idx}. {claim}")
    add_heading(doc, "事实与理由" if route == "civil" else "事实经过")
    for idx, fact in enumerate(data.get("facts") or [], 1):
        p = doc.add_paragraph()
        set_run_font(p.add_run(f"{idx}. {text(fact.get('text'))}"))
    transaction = data.get("transaction") or {}
    if transaction:
        add_heading(doc, "交易及损失信息")
        add_kv_table(doc, [
            ("交易标的", text(transaction.get("subject"))),
            ("金额", text(transaction.get("amount"))),
            ("支付时间", text(transaction.get("paid_at"))),
            ("支付渠道", text(transaction.get("payment_channel"))),
            ("订单/交易号", text(transaction.get("order_or_transfer_id"))),
            ("约定", text(transaction.get("agreement"))),
            ("履行情况", text(transaction.get("performance"))),
        ])
    add_heading(doc, "证据概况")
    doc.add_paragraph(f"证据共 {len(data.get('evidence') or [])} 项，详见《证据目录》及《证据材料》。")
    if route == "civil":
        add_heading(doc, "纠纷解决意愿")
        doc.add_paragraph(f"是否考虑先行调解：{text(data.get('mediation_preference'))}")
        p = doc.add_paragraph()
        set_run_font(p.add_run("此致"), size=11)
        p = doc.add_paragraph()
        set_run_font(p.add_run(text((data.get("court") or {}).get("name"))), size=11)
    else:
        add_heading(doc, "请求事项")
        doc.add_paragraph("请依法对上述事实和线索进行审查，并告知后续需要补充的材料。本文不对任何人的刑事责任作出结论。")
    add_signature(doc, data, "具状人" if route == "civil" else "报案人")
    doc.save(path)


def build_evidence_index_docx(data: dict[str, Any], path: Path, draft: bool) -> None:
    doc = init_doc("证据目录", draft)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    headers = ["序号", "证据名称", "来源及形成时间", "对应文件/页码", "证明对象和内容", "原始载体", "完整性说明"]
    evidence = data.get("evidence") or []
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(1.2), Cm(3.1), Cm(3.8), Cm(3.3), Cm(5.4), Cm(3.6), Cm(4.6)]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E7E6E6")
        set_cell_margins(cell, top=80, bottom=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), name="黑体", size=9.5, bold=True)
    set_repeat_table_header(table.rows[0])
    for idx, item in enumerate(evidence, 1):
        row = table.add_row()
        values = [
            str(idx),
            text(item.get("name")),
            f"{text(item.get('source'))}\n{text(item.get('formed_at'), '')}",
            f"{text(item.get('file'), '')}\n{text(item.get('page_or_range'), '')}",
            text(item.get("proof_purpose")),
            text(item.get("original_carrier")),
            text(item.get("integrity_note")),
        ]
        for cidx, value in enumerate(values):
            cell = row.cells[cidx]
            cell.width = widths[cidx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=80, bottom=80)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if cidx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), size=9)
    set_table_borders(table)
    add_signature(doc, data, "提交人")
    doc.save(path)


def build_upload_checklist_docx(data: dict[str, Any], path: Path, draft: bool) -> None:
    doc = init_doc("人民法院在线服务上传清单", draft)
    add_heading(doc, "建议上传顺序")
    for item in ("01-民事起诉状.pdf", "02-证据目录.pdf", "03-证据材料.pdf", "05-原告身份证明.pdf", "平台或法院要求的其他材料"):
        doc.add_paragraph(f"□ {item}")
    add_heading(doc, "提交前确认")
    for item in (
        "已在当前页面核对允许的格式、单文件大小、文件数量和材料类别",
        "起诉状、证据目录和证据材料中的姓名、金额、日期、编号一致",
        "已逐页检查清晰度、方向、完整性及敏感信息",
        "已由本人完成签名或盖章并填写提交日期",
        "已保留所有原始证据及未压缩副本",
    ):
        doc.add_paragraph(f"□ {item}")
    doc.save(path)


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("CJKTitle", parent=styles["Title"], fontName=FONT_CN, fontSize=18, leading=25, alignment=TA_CENTER, textColor=colors.black, spaceAfter=10 * mm),
        "h": ParagraphStyle("CJKHeading", parent=styles["Heading2"], fontName=FONT_CN, fontSize=13, leading=19, textColor=colors.black, spaceBefore=5 * mm, spaceAfter=2 * mm),
        "body": ParagraphStyle("CJKBody", parent=styles["BodyText"], fontName=FONT_CN, fontSize=10.5, leading=17, textColor=colors.black, alignment=TA_LEFT, spaceAfter=2 * mm),
        "small": ParagraphStyle("CJKSmall", parent=styles["BodyText"], fontName=FONT_CN, fontSize=8, leading=11, textColor=colors.black),
    }


def pdf_header_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont(FONT_CN, 8)
    canvas.drawCentredString(canvas._pagesize[0] / 2, 12 * mm, f"第 {doc.page} 页")
    canvas.restoreState()


def build_pdf(story, path: Path, pagesize=A4) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=pagesize, leftMargin=20 * mm, rightMargin=20 * mm, topMargin=20 * mm, bottomMargin=20 * mm)
    doc.build(story, onFirstPage=pdf_header_footer, onLaterPages=pdf_header_footer)


def p(value: Any, style) -> Paragraph:
    escaped = text(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
    return Paragraph(escaped, style)


def build_readiness_pdf(data: dict[str, Any], issues: list[str], path: Path, draft: bool) -> None:
    s = pdf_styles()
    story = [p(("草稿-待核对\n" if draft else "") + "材料准备情况及待确认事项", s["title"]), p("当前建议路径", s["h"])]
    route_labels = {"civil": "优先准备民事起诉材料", "police": "优先准备报案与证据保全材料", "review": "暂缓形成最终文书，先人工核验"}
    story.append(p(route_labels.get(data.get("route"), "待确认"), s["body"]))
    story.append(p("仍需处理", s["h"]))
    for idx, issue in enumerate(issues or ["未发现阻止生成最终版本的关键缺项；仍须由本人逐项核对。"], 1):
        story.append(p(f"{idx}. {issue}", s["body"]))
    story.append(p("重要提示", s["h"]))
    for item in ("本材料不代表法院或公安机关将受理、立案或支持相关请求。", "请保留原始设备、原始文件及平台可下载记录。", "正式提交前请本人逐项核对并签名。"):
        story.append(p(item, s["body"]))
    build_pdf(story, path)


def build_primary_pdf(data: dict[str, Any], path: Path, draft: bool) -> None:
    s = pdf_styles()
    # Keep a typical one-transaction small-claim complaint on one page while
    # preserving readable spacing and a real handwritten-signature area.
    s["h"].spaceBefore = 3 * mm
    s["h"].spaceAfter = 1 * mm
    s["body"].leading = 15
    s["body"].spaceAfter = 1 * mm
    route = data.get("route")
    title = "民事起诉状" if route == "civil" else "报案事实说明"
    story = [p(("草稿-待核对\n" if draft else "") + title, s["title"]), p("当事人信息", s["h"])]
    plaintiff = data.get("plaintiff") or {}
    rows = [[p("原告" if route == "civil" else "报案人", s["body"]), p(person_summary(plaintiff), s["body"])]]
    for idx, defendant in enumerate(data.get("defendants") or [], 1):
        label = "被告" if route == "civil" and idx == 1 else f"被告{idx}" if route == "civil" else f"对方/线索{idx}"
        rows.append([p(label, s["body"]), p(person_summary(defendant), s["body"])])
    table = Table(rows, colWidths=[35 * mm, 125 * mm], repeatRows=0)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F2F2")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
    story.append(table)
    if route == "civil":
        court = data.get("court") or {}
        story += [p("受诉法院及案由", s["h"]), p(f"受诉法院：{text(court.get('name'))}", s["body"]), p(f"管辖核验：{text(court.get('basis'))}", s["body"]), p(f"案由建议：{text(data.get('title'))}", s["body"]), p("诉讼请求", s["h"])]
        for idx, claim in enumerate(data.get("claims") or [], 1):
            story.append(p(f"{idx}. {claim}", s["body"]))
    story.append(p("事实与理由" if route == "civil" else "事实经过", s["h"]))
    for idx, fact in enumerate(data.get("facts") or [], 1):
        story.append(p(f"{idx}. {text(fact.get('text'))}", s["body"]))
    transaction = data.get("transaction") or {}
    if transaction:
        story.append(p("交易及损失信息", s["h"]))
        for label, key in (("交易标的", "subject"), ("金额", "amount"), ("支付时间", "paid_at"), ("支付渠道", "payment_channel"), ("订单/交易号", "order_or_transfer_id"), ("约定", "agreement"), ("履行情况", "performance")):
            story.append(p(f"{label}：{text(transaction.get(key))}", s["body"]))
    story += [p("证据概况", s["h"]), p(f"证据共 {len(data.get('evidence') or [])} 项，详见《证据目录》及《证据材料》。", s["body"])]
    if route == "civil":
        story += [p("纠纷解决意愿", s["h"]), p(f"是否考虑先行调解：{text(data.get('mediation_preference'))}", s["body"])]
        story += [p("此致", s["body"]), p(text((data.get("court") or {}).get("name")), s["body"])]
        role = "具状人"
    else:
        story += [p("请求事项", s["h"]), p("请依法对上述事实和线索进行审查，并告知后续需要补充的材料。本文不对任何人的刑事责任作出结论。", s["body"])]
        role = "报案人"
    signature = data.get("signature") or {}
    story += [p(f"{role}姓名（供核对）：{text(signature.get('name'))}", s["body"]), p("本人签名/盖章：________________", s["body"]), p(f"日期：{text(signature.get('date'))}", s["body"])]
    build_pdf(story, path)


def build_evidence_index_pdf(data: dict[str, Any], path: Path, draft: bool) -> None:
    s = pdf_styles()
    story = [p(("草稿-待核对\n" if draft else "") + "证据目录", s["title"])]
    headers = ["序号", "证据名称", "来源及形成时间", "对应文件/页码", "证明对象和内容", "原始载体", "完整性说明"]
    rows = [[p(h, s["small"]) for h in headers]]
    for idx, item in enumerate(data.get("evidence") or [], 1):
        rows.append([
            p(str(idx), s["small"]), p(item.get("name"), s["small"]), p(f"{text(item.get('source'))}\n{text(item.get('formed_at'), '')}", s["small"]),
            p(f"{text(item.get('file'), '')}\n{text(item.get('page_or_range'), '')}", s["small"]), p(item.get("proof_purpose"), s["small"]), p(item.get("original_carrier"), s["small"]), p(item.get("integrity_note"), s["small"]),
        ])
    table = LongTable(rows, colWidths=[12 * mm, 32 * mm, 39 * mm, 32 * mm, 56 * mm, 36 * mm, 50 * mm], repeatRows=1)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.35, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E7E6E6")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("ALIGN", (0, 0), (0, -1), "CENTER"), ("LEFTPADDING", (0, 0), (-1, -1), 2), ("RIGHTPADDING", (0, 0), (-1, -1), 2), ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]))
    story.append(table)
    signature = data.get("signature") or {}
    story += [Spacer(1, 4 * mm), p(f"提交人姓名（供核对）：{text(signature.get('name'))}", s["body"]), p("本人签名/盖章：________________", s["body"]), p(f"日期：{text(signature.get('date'))}", s["body"])]
    build_pdf(story, path, pagesize=landscape(A4))


def build_upload_checklist_pdf(data: dict[str, Any], path: Path, draft: bool) -> None:
    s = pdf_styles()
    story = [p(("草稿-待核对\n" if draft else "") + "人民法院在线服务上传清单", s["title"]), p("建议上传顺序", s["h"])]
    for item in ("01-民事起诉状.pdf", "02-证据目录.pdf", "03-证据材料.pdf", "05-原告身份证明.pdf", "平台或法院要求的其他材料"):
        story.append(p(f"□ {item}", s["body"]))
    story.append(p("提交前确认", s["h"]))
    for item in ("已核对当前页面的格式、大小、数量和材料类别", "三份材料中的姓名、金额、日期和编号一致", "已逐页检查清晰度、方向和完整性", "已由本人签名或盖章并填写日期", "已保留所有原始证据"):
        story.append(p(f"□ {item}", s["body"]))
    build_pdf(story, path)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Convert with LibreOffice so DOCX and PDF have identical content/layout."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice/soffice is required to create matching PDF files")
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="court-packet-lo-") as profile:
        profile_uri = Path(profile).resolve().as_uri()
        env = os.environ.copy()
        env.setdefault("TMPDIR", "/private/tmp")
        env.setdefault(
            "SAL_FONTPATH",
            "/System/Library/Fonts:/System/Library/Fonts/Supplemental:/Library/Fonts",
        )
        proc = subprocess.run(
            [
                soffice,
                "--headless",
                f"-env:UserInstallation={profile_uri}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(pdf_path.parent),
                str(docx_path),
            ],
            check=False,
            capture_output=True,
            text=True,
            env=env,
        )
        generated = pdf_path.parent / f"{docx_path.stem}.pdf"
        if proc.returncode != 0 or not generated.exists() or generated.stat().st_size == 0:
            raise RuntimeError(f"DOCX to PDF conversion failed: {proc.stderr or proc.stdout}")
        if generated.resolve() != pdf_path.resolve():
            shutil.move(str(generated), str(pdf_path))


def make_evidence_cover(item: dict[str, Any], number: int, path: Path) -> None:
    s = pdf_styles()
    story = [
        p(f"证据{number}", s["title"]),
        p(text(item.get("name")), s["h"]),
        p(f"来源及形成时间：{text(item.get('source'))}；{text(item.get('formed_at'), '')}", s["body"]),
        p(f"证明对象和内容：{text(item.get('proof_purpose'))}", s["body"]),
        p(f"原始载体：{text(item.get('original_carrier'))}", s["body"]),
        p(f"完整性说明：{text(item.get('integrity_note'))}", s["body"]),
    ]
    build_pdf(story, path)


def image_to_pdf(image_path: Path, pdf_path: Path) -> None:
    with Image.open(image_path) as im:
        fixed = ImageOps.exif_transpose(im).convert("RGB")
        page_w, page_h = 1240, 1754
        margin = 90
        scale = min((page_w - 2 * margin) / fixed.width, (page_h - 2 * margin) / fixed.height)
        resized = fixed.resize((max(1, int(fixed.width * scale)), max(1, int(fixed.height * scale))), Image.Resampling.LANCZOS)
        page = Image.new("RGB", (page_w, page_h), "white")
        page.paste(resized, ((page_w - resized.width) // 2, (page_h - resized.height) // 2))
        page.save(pdf_path, "PDF", resolution=150.0)


def build_evidence_packet(data: dict[str, Any], path: Path) -> None:
    writer = PdfWriter()
    base = Path(data.get("_base_dir", "."))
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        for idx, item in enumerate(data.get("evidence") or [], 1):
            cover = tmpdir_path / f"cover-{idx}.pdf"
            make_evidence_cover(item, idx, cover)
            for page in PdfReader(str(cover)).pages:
                writer.add_page(page)
            raw_files = item.get("files") or ([item.get("file")] if item.get("file") else [])
            for file_idx, raw_file in enumerate(raw_files, 1):
                source = Path(raw_file)
                if not source.is_absolute():
                    source = base / source
                if not source.exists():
                    continue
                suffix = source.suffix.lower()
                if suffix == ".pdf":
                    for page in PdfReader(str(source)).pages:
                        writer.add_page(page)
                elif suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp"}:
                    converted = tmpdir_path / f"image-{idx}-{file_idx}.pdf"
                    image_to_pdf(source, converted)
                    for page in PdfReader(str(converted)).pages:
                        writer.add_page(page)
        with path.open("wb") as f:
            writer.write(f)


def build_identity_packet(data: dict[str, Any], path: Path) -> None:
    documents = data.get("identity_documents") or []
    if not documents:
        return
    base = Path(data.get("_base_dir", "."))
    writer = PdfWriter()
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        page_no = 0
        for document in documents:
            raw_files = document.get("rendered_files") or document.get("files") or []
            for raw_file in raw_files:
                source = Path(raw_file)
                if not source.is_absolute():
                    source = base / source
                if not source.exists() or source.suffix.lower() not in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp"}:
                    continue
                page_no += 1
                converted = tmpdir_path / f"identity-{page_no}.pdf"
                image_to_pdf(source, converted)
                for page in PdfReader(str(converted)).pages:
                    writer.add_page(page)
        if page_no:
            with path.open("wb") as f:
                writer.write(f)


def main() -> int:
    args = parse_args()
    data = load_case(args.case_json)
    issues = validate(data, args.mode)
    out = args.output_dir
    out.mkdir(parents=True, exist_ok=True)
    draft = args.mode == "draft"
    suffix = "（草稿-待核对）" if draft else ""
    primary = "民事起诉状" if data.get("route") == "civil" else "报案事实说明"

    readiness_docx = out / f"00-材料准备情况及待确认事项{suffix}.docx"
    readiness_pdf = out / f"00-材料准备情况及待确认事项{suffix}.pdf"
    build_readiness_docx(data, issues, readiness_docx, draft)
    build_readiness_pdf(data, issues, readiness_pdf, draft)
    if data.get("route") != "review":
        primary_docx = out / f"01-{primary}{suffix}.docx"
        primary_pdf = out / f"01-{primary}{suffix}.pdf"
        build_primary_docx(data, primary_docx, draft)
        build_primary_pdf(data, primary_pdf, draft)
    index_docx = out / f"02-证据目录{suffix}.docx"
    index_pdf = out / f"02-证据目录{suffix}.pdf"
    build_evidence_index_docx(data, index_docx, draft)
    build_evidence_index_pdf(data, index_pdf, draft)
    build_evidence_packet(data, out / f"03-证据材料{suffix}.pdf")
    if data.get("route") == "civil":
        checklist_docx = out / f"04-人民法院在线服务上传清单{suffix}.docx"
        checklist_pdf = out / f"04-人民法院在线服务上传清单{suffix}.pdf"
        build_upload_checklist_docx(data, checklist_docx, draft)
        build_upload_checklist_pdf(data, checklist_pdf, draft)
        build_identity_packet(data, out / f"05-原告身份证明{suffix}.pdf")
    print(json.dumps({"output_dir": str(out.resolve()), "mode": args.mode, "blocking_issues": issues}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(2)
